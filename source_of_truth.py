"""Source of Truth — GitHub-backed wiki + how-to knowledge base."""

import base64
import json
import re
import threading
import uuid
import time

import psycopg2
import psycopg2.extras
from flask import Blueprint, render_template, request, jsonify, redirect, url_for
from openai import OpenAI

from _sot_github import (
    save_howto, delete_howto_file,
    read_page, list_pages, is_configured, rebuild_index_from_repo,
    append_insights, read_insights_raw, INSIGHTS_PATH,
    append_concepts, read_concepts_raw, CONCEPTS_PATH,
    append_products, read_products_raw, PRODUCTS_PATH,
)
from _sot_dedup import (
    find_lookalikes, load_existing,
    merge_source_into_insight, merge_source_into_product,
)

import os

NAME = "Source of Truth"
blueprint = Blueprint("source_of_truth", __name__, template_folder="templates/source_of_truth")


@blueprint.context_processor
def _inject_repo():
    return {"repo": os.environ.get("SOT_GITHUB_REPO", "")}

llm = OpenAI(
    base_url=os.environ.get("OPENAI_BASE_URL", "https://api.openai.com/v1"),
    api_key=os.environ.get("OPENAI_API_KEY", "unused"),
)
LLM_MODEL = os.environ.get("SOT_EXTRACT_MODEL", "anthropic/claude-sonnet-4.5")

_jobs = {}
_jobs_lock = threading.Lock()


def get_db():
    """DB only used for pending extractions (in-flight LLM work).

    Reads DATABASE_URL (psycopg2 conninfo or postgres:// URL).
    """
    dsn = os.environ.get("DATABASE_URL")
    if not dsn:
        raise RuntimeError("DATABASE_URL is not set. See .env.example.")
    return psycopg2.connect(dsn)


def slugify(text):
    return re.sub(r'[^a-z0-9]+', '-', text.lower().strip()).strip('-')[:80]


_HOWTO_PATH_RE = re.compile(r"^howtos/[a-z0-9][a-z0-9-]{0,99}\.md$")


def _validate_howto_path(path):
    """Return True iff `path` looks like a legitimate how-to file path.

    Defends against the `<path:path>` route variable being abused to read or
    delete arbitrary files in the SOT GitHub repo. We accept only paths of
    the shape `howtos/<slug>.md` where slug matches our own slugify output.
    """
    if not path or len(path) > 120:
        return False
    if ".." in path or path.startswith("/"):
        return False
    return bool(_HOWTO_PATH_RE.match(path))


def _parse_page_content(raw):
    """Strip the '# Title' and '> Category:' header from stored markdown, return (title, category, body)."""
    lines = (raw or "").split("\n")
    title, category, body_start = "", "", 0
    for i, line in enumerate(lines):
        if line.startswith("# ") and not title:
            title = line[2:].strip()
        elif line.startswith("> Category:"):
            category = line.split(":", 1)[1].strip()
        elif title and (line.strip() == "" or line.startswith("> ")):
            continue
        else:
            body_start = i
            break
    body = "\n".join(lines[body_start:]).strip()
    return title, category, body


def ensure_schema():
    """Create both Postgres tables this app needs (idempotent)."""
    from _sot_dedup import ensure_schema as _dedup_schema
    _dedup_schema()
    with get_db() as conn, conn.cursor() as cur:
        cur.execute("""
            CREATE TABLE IF NOT EXISTS sot_pending_updates (
                id                SERIAL PRIMARY KEY,
                source_type       TEXT NOT NULL,
                source_ref        TEXT,
                source_content    TEXT,
                status            TEXT NOT NULL DEFAULT 'pending',
                extracted_facts   JSONB,
                proposed_changes  JSONB,
                created_at        TIMESTAMPTZ NOT NULL DEFAULT now()
            )
        """)
        cur.execute("""
            CREATE INDEX IF NOT EXISTS sot_pending_updates_status_idx
                ON sot_pending_updates (status, created_at DESC)
        """)
        conn.commit()


# ── Routes ──────────────────────────────────────────────────────

@blueprint.route("/")
def dashboard():
    index = list_pages()
    insights_entry = next((e for e in index if e.get("type") == "insights"), None)
    concepts_entry = next((e for e in index if e.get("type") == "concepts"), None)
    products_entry = next((e for e in index if e.get("type") == "products"), None)
    howtos = [e for e in index if e.get("type") == "howto"]

    # Recent insight + concept + product cards (newest first, top 5 each)
    recent_insights, insights_count = [], 0
    if insights_entry:
        cards = _split_insight_cards(read_insights_raw())
        insights_count = len(cards)
        recent_insights = cards[-5:][::-1]

    recent_concepts, concepts_count = [], 0
    if concepts_entry:
        cards = _split_insight_cards(read_concepts_raw())
        concepts_count = len(cards)
        recent_concepts = cards[-5:][::-1]

    recent_products, products_count = [], 0
    if products_entry:
        cards = _split_insight_cards(read_products_raw())
        products_count = len(cards)
        recent_products = cards[-5:][::-1]

    conn = get_db()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    cur.execute("SELECT * FROM sot_pending_updates WHERE status IN ('pending','extracting') ORDER BY created_at DESC")
    pending = cur.fetchall()
    conn.close()

    return render_template(
        "sot_dashboard.html",
        howtos=howtos,
        insights_count=insights_count,
        recent_insights=recent_insights,
        concepts_count=concepts_count,
        recent_concepts=recent_concepts,
        products_count=products_count,
        recent_products=recent_products,
        pending=pending,
    )


def _split_insight_cards(raw):
    """Split insights.md body into a list of card markdown blocks (one per ### heading)."""
    if not raw:
        return []
    # Strip the top-of-file header before the first ###
    parts = raw.split("\n### ")
    if len(parts) <= 1:
        return []
    return ["### " + p.strip() for p in parts[1:]]

def _json_value(value, default=None):
    """Return jsonb values as Python objects whether psycopg2 gave us dict/list or str."""
    if value is None:
        return default
    if isinstance(value, (dict, list)):
        return value
    try:
        return json.loads(value)
    except Exception:
        return default


def _source_key(value):
    raw = (value or "").encode("utf-8")
    return base64.urlsafe_b64encode(raw).decode("ascii").rstrip("=")


def _source_from_key(key):
    try:
        padded = key + "=" * (-len(key) % 4)
        return base64.urlsafe_b64decode(padded.encode("ascii")).decode("utf-8")
    except Exception:
        return ""


def _card_title(card_md):
    first = (card_md or "").split("\n", 1)[0]
    return first.replace("### ", "").strip()


def _extract_card_meta(card_md):
    """Extract lightweight metadata from one rendered SOT markdown card."""
    meta = {"title": _card_title(card_md), "markdown": card_md}
    m = re.search(r"- \*\*Source:\*\* \[([^\]]+)\]\(([^)]+)\)", card_md or "")
    if m:
        meta["source_title"] = m.group(1).strip()
        meta["source_url"] = m.group(2).strip()
    else:
        m = re.search(r"- \*\*Source:\*\*\s*(.+)", card_md or "")
        if m:
            meta["source_title"] = m.group(1).strip()
            meta["source_url"] = meta["source_title"]
    m = re.search(r"- \*\*Date:\*\*\s*(.+)", card_md or "")
    if m:
        meta["date"] = m.group(1).strip()
    return meta


def _current_sot_cards_by_source():
    """Group cards currently saved in the GitHub SOT by source URL/title."""
    grouped = {}
    for bucket, raw in [
        ("insights", read_insights_raw()),
        ("concepts", read_concepts_raw()),
        ("products", read_products_raw()),
    ]:
        for card_md in _split_insight_cards(raw):
            meta = _extract_card_meta(card_md)
            key = meta.get("source_url") or meta.get("source_title") or ""
            if not key:
                continue
            grouped.setdefault(key, {"insights": [], "concepts": [], "products": [], "howtos": []})[bucket].append(meta)
    return grouped


def _items_by_source(proposed):
    """Group proposed extracted items by source_url/source_title for source detail pages."""
    grouped = {}
    for bucket in ("insights", "concepts", "products", "howtos"):
        for item in (proposed.get(bucket) or []):
            key = (item.get("source_url") or item.get("source_title") or "Unknown source").strip()
            grouped.setdefault(key, {"insights": [], "concepts": [], "products": [], "howtos": []})[bucket].append(item)
    return grouped


def _item_title(bucket, item):
    if bucket == "concepts":
        return item.get("concept") or "Untitled concept"
    return item.get("claim") or item.get("title") or "Untitled item"


def _source_title_from_items(key, buckets):
    for bucket_items in buckets.values():
        for item in bucket_items:
            if item.get("source_title"):
                return item.get("source_title")
    return key


def _source_rows_from_pending(pu):
    """Expand one pending-update row into one or more human-facing source rows."""
    proposed = _json_value(pu.get("proposed_changes"), {}) or {}
    extracted = _json_value(pu.get("extracted_facts"), {}) or {}
    by_source = _items_by_source(proposed)
    rows = {}

    def ensure(key, title=None, source_type=None, status=None, error=None):
        if not key:
            key = pu.get("source_ref") or pu.get("source_content") or f"pending:{pu['id']}"
        if key not in rows:
            rows[key] = {
                "pending_id": pu["id"],
                "source_key": _source_key(key),
                "source_id": key,
                "title": title or key,
                "url": key if str(key).startswith(("http://", "https://")) else "",
                "source_type": source_type or pu.get("source_type") or "unknown",
                "status": status or pu.get("status") or "unknown",
                "created_at": pu.get("created_at"),
                "last_action": pu.get("status") or "unknown",
                "error": error or "",
                "counts": {"insights": 0, "concepts": 0, "products": 0, "howtos": 0},
            }
        return rows[key]

    # Prefer explicit per-URL / per-file status entries when available.
    for entry in proposed.get("per_url") or []:
        key = entry.get("url") or "Unknown URL"
        row = ensure(key, key, "url", entry.get("status") or pu.get("status"), entry.get("error"))
        row["counts"] = {
            "insights": entry.get("insight_count", 0),
            "concepts": entry.get("concept_count", 0),
            "products": entry.get("product_count", 0),
            "howtos": entry.get("howto_count", 0),
        }
    for entry in proposed.get("per_file") or []:
        key = entry.get("name") or "Unknown file"
        row = ensure(key, key, "file", entry.get("status") or pu.get("status"), entry.get("error"))
        row["counts"] = {
            "insights": entry.get("insight_count", 0),
            "concepts": entry.get("concept_count", 0),
            "products": entry.get("product_count", 0),
            "howtos": entry.get("howto_count", 0),
        }

    # Add any source URLs found directly on extracted items.
    for key, buckets in by_source.items():
        row = ensure(key, _source_title_from_items(key, buckets), pu.get("source_type"))
        for bucket, items in buckets.items():
            row["counts"][bucket] = max(row["counts"].get(bucket, 0), len(items))

    # Fallback for plain text / old rows without proposed source grouping.
    if not rows:
        key = pu.get("source_ref") or pu.get("source_content") or f"pending:{pu['id']}"
        row = ensure(key, key, pu.get("source_type"))
        row["counts"] = {
            "insights": extracted.get("insights_count", 0),
            "concepts": extracted.get("concepts_count", 0),
            "products": extracted.get("products_count", 0),
            "howtos": extracted.get("howtos_count", 0),
        }

    return list(rows.values())


def _all_source_rows(q="", source_type="", status=""):
    conn = get_db()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    cur.execute("SELECT * FROM sot_pending_updates ORDER BY created_at DESC, id DESC")
    pending_rows = cur.fetchall()
    conn.close()
    rows = []
    for pu in pending_rows:
        rows.extend(_source_rows_from_pending(pu))
    # Attach saved-card counts from the actual SOT files.
    current = _current_sot_cards_by_source()
    for row in rows:
        saved = current.get(row["source_id"], {})
        row["saved_count"] = sum(len(saved.get(b, [])) for b in ("insights", "concepts", "products", "howtos"))
    q = (q or "").strip().lower()
    if q:
        rows = [r for r in rows if q in (r.get("title", "") + " " + r.get("source_id", "")).lower()]
    if source_type:
        rows = [r for r in rows if r.get("source_type") == source_type]
    if status:
        rows = [r for r in rows if r.get("status") == status or r.get("last_action") == status]
    return rows


@blueprint.route("/sources")
def sources_view():
    q = (request.args.get("q") or "").strip()
    source_type = (request.args.get("type") or "").strip()
    status = (request.args.get("status") or "").strip()
    rows = _all_source_rows(q=q, source_type=source_type, status=status)
    type_options = sorted({r.get("source_type") for r in _all_source_rows() if r.get("source_type")})
    status_options = sorted({r.get("status") for r in _all_source_rows() if r.get("status")})
    return render_template(
        "sot_sources.html",
        sources=rows,
        q=q,
        source_type=source_type,
        status=status,
        type_options=type_options,
        status_options=status_options,
    )


@blueprint.route("/sources/<int:pu_id>/<source_key>")
def source_detail(pu_id, source_key):
    source_id = _source_from_key(source_key)
    conn = get_db()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    cur.execute("SELECT * FROM sot_pending_updates WHERE id = %s", (pu_id,))
    pu = cur.fetchone()
    conn.close()
    if not pu:
        return redirect(url_for("source_of_truth.sources_view"))

    rows = _source_rows_from_pending(pu)
    source = next((r for r in rows if r["source_id"] == source_id), None)
    if not source:
        source = {"pending_id": pu_id, "source_key": source_key, "source_id": source_id, "title": source_id, "source_type": pu.get("source_type"), "status": pu.get("status"), "created_at": pu.get("created_at"), "counts": {}}

    proposed = _json_value(pu.get("proposed_changes"), {}) or {}
    proposed_grouped = _items_by_source(proposed).get(source_id, {"insights": [], "concepts": [], "products": [], "howtos": []})

    saved_grouped = _current_sot_cards_by_source().get(source_id, {"insights": [], "concepts": [], "products": [], "howtos": []})

    # How-to files don't currently store source metadata. For applied rows, infer saved how-tos
    # by matching proposed how-to titles against current how-to index entries.
    howto_titles = {h.get("title") for h in proposed_grouped.get("howtos", []) if h.get("title")}
    if howto_titles:
        for h in [e for e in list_pages() if e.get("type") == "howto"]:
            if h.get("title") in howto_titles:
                saved_grouped.setdefault("howtos", []).append({"title": h.get("title"), "path": h.get("path"), "summary": h.get("summary", "")})

    return render_template(
        "sot_source_detail.html",
        source=source,
        pu=pu,
        saved=saved_grouped,
        proposed=proposed_grouped,
        item_title=_item_title,
    )


# ── Insights view ───────────────────────────────────────────────

@blueprint.route("/insights")
def insights_view():
    raw = read_insights_raw()
    cards = _split_insight_cards(raw)
    return render_template("sot_insights.html", cards=cards[::-1], total=len(cards),
                           title="Insights & Stats", description="quotable fact", icon="📊")


@blueprint.route("/products")
def products_view():
    raw = read_products_raw()
    cards = _split_insight_cards(raw)
    return render_template("sot_insights.html", cards=cards[::-1], total=len(cards),
                           title="Product Info", description="product fact", icon="📦")


@blueprint.route("/concepts")
def concepts_view():
    raw = read_concepts_raw()
    cards = _split_insight_cards(raw)
    return render_template("sot_insights.html", cards=cards[::-1], total=len(cards),
                           title="Concepts & Mechanisms", description="explainer", icon="🧩")


# ── How-Tos ─────────────────────────────────────────────────────

@blueprint.route("/howtos")
def howto_list():
    index = list_pages()
    howtos = [e for e in index if e["type"] == "howto"]
    return render_template("sot_howto_list.html", howtos=howtos)


@blueprint.route("/howto/new", methods=["GET", "POST"])
def howto_new():
    if request.method == "POST":
        data = request.get_json()
        title = data.get("title", "").strip()
        if not title:
            return jsonify({"error": "Title required"}), 400
        path, slug = save_howto(title, data.get("description", ""), data.get("prerequisites", ""),
                                data.get("steps", []), data.get("tags", []))
        time.sleep(2)
        return jsonify({"path": path, "slug": slug})
    return render_template("sot_howto_edit.html", howto=None, is_new=True)


def _parse_howto_content(raw):
    """Parse a how-to markdown file into structured data."""
    lines = (raw or "").split("\n")
    title, desc, prereqs, tags = "", "", "", []
    steps = []
    section = "desc"
    for line in lines:
        if line.startswith("# "):
            title = line[2:].strip()
        elif line.startswith("## Prerequisites"):
            section = "prereqs"
        elif line.startswith("## Steps"):
            section = "steps"
        elif line.startswith("**Tags:**"):
            tags = [t.strip() for t in line.replace("**Tags:**", "").split(",") if t.strip()]
        elif section == "prereqs" and line.strip():
            prereqs += line.strip() + " "
        elif section == "steps" and re.match(r'^\d+\.', line.strip()):
            step_text = re.sub(r'^\d+\.\s*', '', line.strip())
            # Parse bold title
            m = re.match(r'\*\*(.+?)\*\*(.*)', step_text)
            if m:
                steps.append({"title": m.group(1), "description": m.group(2).strip()})
            else:
                steps.append({"title": step_text, "description": ""})
        elif section == "steps" and line.strip().startswith("   ") and steps:
            steps[-1]["description"] = (steps[-1].get("description", "") + " " + line.strip()).strip()
        elif section == "desc" and line.strip() and not line.startswith(">"):
            desc += line.strip() + " "

    return {
        "title": title,
        "description": desc.strip(),
        "prerequisites": prereqs.strip(),
        "steps": steps,
        "tags": tags,
    }


@blueprint.route("/howto/view/<path:path>")
def howto_view(path):
    if not _validate_howto_path(path):
        return redirect(url_for("source_of_truth.howto_list"))
    raw = read_page(path)
    if not raw:
        return redirect(url_for("source_of_truth.howto_list"))
    howto = _parse_howto_content(raw)
    howto["path"] = path
    return render_template("sot_howto_view.html", howto=howto)


@blueprint.route("/howto/edit/<path:path>", methods=["GET", "POST"])
def howto_edit(path):
    if not _validate_howto_path(path):
        return redirect(url_for("source_of_truth.howto_list"))
    raw = read_page(path)
    if not raw:
        return redirect(url_for("source_of_truth.howto_list"))
    howto = _parse_howto_content(raw)
    howto["path"] = path

    if request.method == "POST":
        data = request.get_json()
        new_title = data.get("title", howto["title"])
        # Delete old if title changed
        old_slug = slugify(howto["title"])
        new_slug = slugify(new_title)
        if old_slug != new_slug:
            delete_howto_file(howto["title"])
            time.sleep(1)
        save_howto(new_title, data.get("description", ""), data.get("prerequisites", ""),
                   data.get("steps", []), data.get("tags", []))
        time.sleep(2)
        return jsonify({"ok": True, "path": f"howtos/{new_slug}.md"})
    return render_template("sot_howto_edit.html", howto=howto, is_new=False)


@blueprint.route("/howto/delete/<path:path>", methods=["POST"])
def howto_delete(path):
    if not _validate_howto_path(path):
        return redirect(url_for("source_of_truth.howto_list"))
    raw = read_page(path)
    if raw:
        howto = _parse_howto_content(raw)
        delete_howto_file(howto["title"])
    return redirect(url_for("source_of_truth.howto_list"))


# ── Rebuild index ──────────────────────────────────────────────

# ── Search ──────────────────────────────────────────────────────
@blueprint.route("/search")
def search_page():
    q = request.args.get("q", "").strip()
    return render_template("sot_search.html", q=q)


@blueprint.route("/api/search")
def api_search():
    """Simple substring search over titles/summaries from the index.

    The full hybrid (Postgres FTS + embeddings) search used internally
    requires extra infrastructure and is intentionally out of scope here.
    """
    q = request.args.get("q", "").strip().lower()
    limit = min(int(request.args.get("limit", 15)), 30)
    if not q:
        return jsonify({"hits": [], "query": ""})
    index = list_pages()
    hits = []
    for entry in index:
        hay = " ".join([
            str(entry.get("title", "")),
            str(entry.get("summary", "")),
            str(entry.get("category", "")),
            str(entry.get("path", "")),
        ]).lower()
        if q in hay:
            hits.append({
                "title": entry.get("title", ""),
                "snippet": entry.get("summary", ""),
                "url": "/" + entry.get("type", "page") + "/" + (entry.get("path", "") or ""),
                "source": "sot",
                "section": entry.get("category", ""),
            })
        if len(hits) >= limit:
            break
    return jsonify({"hits": hits, "query": q})


@blueprint.route("/api/rebuild-index", methods=["POST"])
def api_rebuild_index():
    """Scan GitHub repo and rebuild index.json from scratch."""
    job_id = str(uuid.uuid4())
    with _jobs_lock:
        _jobs[job_id] = {"status": "running"}

    def _do():
        try:
            index = rebuild_index_from_repo()
            with _jobs_lock:
                _jobs[job_id] = {"status": "done", "count": len(index)}
        except Exception as e:
            with _jobs_lock:
                _jobs[job_id] = {"status": "error", "error": str(e)}

    threading.Thread(target=_do, daemon=True).start()
    return jsonify({"job_id": job_id})


@blueprint.route("/api/job/<job_id>")
def job_status(job_id):
    with _jobs_lock:
        job = _jobs.get(job_id)
    if not job:
        return jsonify({"error": "Not found"}), 404
    return jsonify(job)


# ── Ingest + LLM extraction ────────────────────────────────────

EXTRACT_SYSTEM = """You are an Ahrefs knowledge curator. Read the source content and extract ONLY three kinds of items:

1. AHREFS HOW-TOS — step-by-step procedures that explicitly use an Ahrefs product or feature.
   - STRICT FILTER 1 (product mention): step 1 MUST name an Ahrefs product or feature (Site Explorer, Keywords Explorer, Site Audit, Rank Tracker, Brand Radar, AI Content Helper, Content Explorer, Web Analytics, Top Pages report, Best by Links, etc.).
   - STRICT FILTER 2 (procedural source): the SOURCE TEXT must actually contain a procedure — a numbered list, sequential 'first/then/next/finally', a labeled 'Step 1/Step 2', or a screenshot walkthrough described in text. A NAME-DROP IS NOT A HOW-TO.
   - STRICT FILTER 3 (extractive, not generative): each step must be supported by a quote from the source text. If the article only says "you can find decaying content in Site Explorer" without explaining HOW, DO NOT invent the steps from your background knowledge. Return zero how-tos for that chunk.
   - If the source describes a generic SEO process that DOESN'T name an Ahrefs tool, SKIP IT. Do NOT re-frame.

   NEGATIVE EXAMPLE — do NOT extract a how-to from a passage like this:
   "You can automatically brainstorm seed keywords in Keywords Explorer; summarise top-ranking content in AI Content Helper; analyse any website's best-performing pages in Site Explorer."
   This is a feature MENTION, not a procedure. There are no steps to extract. Return [] for howtos.

   POSITIVE EXAMPLE — a passage like this IS a how-to:
   "To find a competitor's top pages: 1) Open Site Explorer and enter the competitor's domain. 2) Click the Top Pages report in the left sidebar. 3) Sort by Traffic descending. 4) Filter Position 1-10 to focus on pages they rank well for."

   Per-step REQUIRED field `source_quote`: a SHORT verbatim substring (≤3–10 words) of the source text that justifies this step. The substring will be checked for an exact match against the source. If you cannot quote the source for a step, omit that step. If you cannot quote any step, omit the entire how-to.

2. INSIGHTS & STATS — specific, quotable facts. Each must be:
   - Standalone-readable: someone could paste it into a blog post and it makes sense.
   - Specific: a number, %, named study, or concrete finding. Not vibes, not generic advice.
   - From Ahrefs data/research/blog OR a clearly cited external stat that supports an Ahrefs claim.
   - One claim per entry. Do NOT bundle multiple stats into one card.

3. CONCEPTS & MECHANISMS — 'how X works' / 'what Y means' explainers. Each must be:
   - Relevant to Ahrefs use cases: AI search mechanics, ranking signals, crawling/indexing, SEO concepts, LLM behaviour that affects visibility, search-engine architecture, etc.
   - NOT generic marketing theory, NOT business strategy, NOT product release notes.
   - One concept per card. If the article explains multiple distinct concepts, emit multiple cards.
   - Summary must be plain-English, one sentence, understandable without the source.
   - Key points: 3-7 bullets covering the mechanism + practical implications. Each bullet under 25 words.

4. PRODUCT INFO — facts about a specific Ahrefs product or feature, OR a use case (job the product solves). Each must be:
   - About a named Ahrefs product (Site Explorer, Keywords Explorer, Site Audit, Rank Tracker, Brand Radar, AI Content Helper, Content Explorer, Web Analytics, Agent A, Ahrefs API, Ahrefs Evolve, etc.) OR a named sub-feature (Top Pages report, Best by Links, AI Overviews tracking, Custom Prompts, Connectors framework, etc.).
   - One fact OR one use case per card.
   - NOT a generic SEO stat from research (those go to INSIGHTS).
   - NOT a step-by-step procedure (those go to HOWTOS).
   - NOT a mechanism explainer about how search/AI works (those go to CONCEPTS).

   IMPORTANT: numbers are NOT required. A product fact does not need a percentage or quantity. "Site Explorer surfaces pages with declining traffic" is a complete fact. Don't skip product info just because no number is mentioned. Don't try to make insights out of these either — keep them in products.

   Schema fields:
   - `claim` (required): one-sentence statement — becomes the card heading.
   - `product` (required): canonical Ahrefs product/feature name.
   - `fact_type` (required): one of: capability, limit, pricing, integration, plan, api, behaviour, release, use_case, scale.
   - `detail` (optional): one-line expansion, numeric specifics, persona, or example.
   - `source_url`, `source_title`, `date`, `topic_tag` (optional).

   USE CASES specifically (fact_type = "use_case"):
   - Phrase the claim as a JOB the product helps the user do. Verb-first if possible.
   - Format: "<verb> <object> with <product/feature>" or "<product> helps <persona> do <job>".
   - The use case must be supported by the source: the source must describe the product doing this job, not just hint at it.
   - If the source lists multiple use cases for the same product, emit multiple cards (one each).

   POSITIVE EXAMPLES of product facts:
   - claim: "Brand Radar custom prompts cost ~6 checks per prompt per cycle.", fact_type: limit
   - claim: "Site Explorer's Top Pages report supports filtering by Declining traffic.", fact_type: capability
   - claim: "Ahrefs API v3 DELETE on brand-radar-prompts returns 405 — prompts can only be removed via the UI.", fact_type: api
   - claim: "Agent A connectors are typed integrations grouped by provider.", fact_type: capability

   SCALE FACTS (fact_type = "scale"):
   - Numbers about the SIZE or THROUGHPUT of Ahrefs' underlying data platform: index size, crawl rate, keyword counts, backlink graph size, refresh cadence, content corpus size.
   - These are PRODUCT INFO, NOT generic SEO insights. Even though they're numeric, they describe what Ahrefs HAS, not what SEO research has found. Route them to PRODUCT INFO, not INSIGHTS.
   - `product` field for scale facts: use the relevant product/component name (e.g. "Ahrefs Index", "AhrefsBot", "Keywords database", "Backlink index", or a specific tool if scoped).

   POSITIVE EXAMPLES of scale facts:
   - claim: "Ahrefs has indexed over 170 trillion pages.", product: "Ahrefs Index", fact_type: scale
   - claim: "AhrefsBot crawls 5 million pages per minute.", product: "AhrefsBot", fact_type: scale
   - claim: "Ahrefs tracks 41.9 billion keywords.", product: "Keywords database", fact_type: scale
   - claim: "Ahrefs has mapped 3.5 trillion external backlinks.", product: "Backlink index", fact_type: scale

   POSITIVE EXAMPLES of use cases:
   - claim: "Find decaying pages on your site with Site Explorer's Top Pages report.", product: "Site Explorer", fact_type: use_case, detail: "Filter by Declining traffic + low KD to surface easy wins."
   - claim: "Track how often AI Overviews cite your brand vs competitors with Brand Radar.", product: "Brand Radar", fact_type: use_case
   - claim: "Run an autonomous Monday-morning report of pages that lost 20%+ traffic with Agent A.", product: "Agent A", fact_type: use_case, detail: "Schedule a workflow that pulls GSC + Site Explorer data and posts to Slack."
   - claim: "Brainstorm content angles for a new keyword using AI Content Helper.", product: "AI Content Helper", fact_type: use_case

   NEGATIVE EXAMPLES (do NOT emit these as products):
   - "68% of pages have no backlinks." → insight (this is a finding ABOUT the web, derived from Ahrefs data, not a fact about Ahrefs the product)
   - "How AI search engines retrieve content." → concept
   - "To find decaying content, open Site Explorer, then click Top Pages, then…" → howto (a step-by-step procedure, not a use case)
   - Generic vague claims like "Agent A is the future of SEO" → skip entirely

   DISAMBIGUATION between SCALE (product) and INSIGHT:
   - "Ahrefs has indexed 170 trillion pages." → SCALE product fact (the size of Ahrefs' index)
   - "68% of pages on the web have no backlinks." → INSIGHT (a finding about the web that Ahrefs measured)
   - If the number describes Ahrefs' own infrastructure/data → product (scale).
   - If the number describes what they found by looking at the web → insight.

Return valid JSON only, no markdown fences. Schema:
{
  "howtos": [
    {
      "title": "How to ...",
      "steps": [
        {"text": "Open Site Explorer and enter your domain.", "source_quote": "Open Site Explorer"},
        {"text": "Go to the Top Pages report.", "source_quote": "Top Pages report"}
      ]
    }
  ],
  "insights": [
    {
      "claim": "One-sentence quotable statement of the fact.",
      "number": "68%",
      "source_url": "https://ahrefs.com/blog/...",
      "source_title": "Ahrefs Study: ...",
      "date": "2024-03",
      "topic_tag": "backlinks",
      "context": "Optional one-line note about sample/methodology if needed."
    }
  ],
  "concepts": [
    {
      "concept": "How AI search engines retrieve content",
      "summary": "AI search engines route a query through an LLM that calls a retrieval layer, then synthesizes an answer from the returned passages.",
      "key_points": [
        "Query is first rephrased by the LLM into one or more sub-queries.",
        "Sub-queries hit a retrieval index (sometimes Bing, sometimes a proprietary crawler).",
        "Top passages are ranked + summarized into the visible answer.",
        "Citations are post-hoc — the LLM picks which retrieved sources to attribute."
      ],
      "source_url": "https://ahrefs.com/seo/how-ai-search-engines-work",
      "source_title": "How AI Search Engines Work",
      "date": "2025-01",
      "topic_tag": "ai-search"
    }
  ],
  "products": [
    {
      "claim": "Brand Radar custom prompts cost ~6 checks per prompt per cycle.",
      "product": "Brand Radar",
      "fact_type": "limit",
      "detail": "Counts against the monthly plan limit + any PAYG balance.",
      "source_url": "https://ahrefs.com/...",
      "source_title": "Brand Radar docs",
      "date": "2026-04",
      "topic_tag": "brand-radar"
    },
    {
      "claim": "Find decaying pages on your site with Site Explorer's Top Pages report.",
      "product": "Site Explorer",
      "fact_type": "use_case",
      "detail": "Filter by Declining traffic + low KD to surface easy wins.",
      "source_url": "https://ahrefs.com/...",
      "topic_tag": "site-explorer"
    }
  ]
}

Rules:
- If nothing qualifies for a bucket, return an empty array for that bucket.
- Never invent numbers, sources, dates, or mechanism details. If a field is unknown, omit it (don't fabricate).
- Prefer the article's own published URL as source_url. If only a URL was provided to you and the content references "this article", use that URL.
- Be aggressive on filtering. Quality over quantity. Better to return 0 items than to lower the bar.
- A single article may yield items in multiple buckets simultaneously. Don't force everything into one bucket."""


@blueprint.route("/ingest", methods=["GET"])
def ingest_form():
    return render_template("sot_ingest.html")


# ── Chunking & extraction config ──────────────────────────────
CHUNK_SIZE = 11000          # chars per chunk (LLM context-friendly)
CHUNK_OVERLAP = 500         # chars of overlap to preserve sentence boundaries across chunks
MAX_CHUNKS_PER_SOURCE = 5   # hard cap to bound LLM cost (~$0.05/chunk → max ~$0.25/source)
FETCH_MAX_LENGTH = 60000    # chars to pull from a URL (web-fetch arg) — ~5 chunks worth


def _chunk_text(text):
    """Split a long string into chunks with overlap. Tries to break on paragraph then sentence boundaries.
    Returns list of (chunk_index, total_chunks, chunk_text). Capped at MAX_CHUNKS_PER_SOURCE."""
    text = text or ""
    if len(text) <= CHUNK_SIZE:
        return [(1, 1, text)]

    chunks = []
    pos = 0
    while pos < len(text) and len(chunks) < MAX_CHUNKS_PER_SOURCE:
        end = pos + CHUNK_SIZE
        if end >= len(text):
            chunks.append(text[pos:])
            break
        # Try to land on a paragraph break, then sentence end, then space, in the last 1500 chars
        window_start = max(pos + CHUNK_SIZE - 1500, pos + 1)
        candidates = [
            text.rfind("\n\n", window_start, end),
            text.rfind(". ", window_start, end),
            text.rfind(" ", window_start, end),
        ]
        cut = max(candidates)
        if cut <= pos:
            cut = end  # fallback hard cut
        chunks.append(text[pos:cut])
        pos = max(cut - CHUNK_OVERLAP, cut)  # rewind for overlap
    total = len(chunks)
    return [(i + 1, total, c) for i, c in enumerate(chunks)]


# ── Robust JSON parsing for LLM extraction output ──────────────────────────

def _parse_extraction_json(raw, chunk_idx=None):
    """Best-effort JSON parsing of LLM extraction output.
    Strategy:
      1. Strip optional markdown fences.
      2. Try strict json.loads.
      3. On failure, try truncating to the last balanced `}` and parsing the prefix
         (handles cases where the model trailed off into garbage).
      4. On failure, return a permissive empty result and log a warning rather than crashing
         the whole ingest. Better to lose one chunk than the whole article.
    """
    if not raw:
        return {"howtos": [], "insights": [], "concepts": [], "products": []}
    s = raw.strip()
    s = re.sub(r'^```(?:json)?\s*', '', s)
    s = re.sub(r'\s*```$', '', s).strip()

    try:
        return json.loads(s)
    except json.JSONDecodeError as e:
        err1 = e

    # Salvage attempt: walk braces and find the last balanced position.
    last_balanced = _find_last_balanced_object_end(s)
    if last_balanced is not None and last_balanced > 0:
        try:
            return json.loads(s[: last_balanced + 1])
        except json.JSONDecodeError:
            pass

    # Final fallback: log + return empty so the rest of the article keeps moving.
    snippet = s[: 200].replace("\n", " ")
    print(f"[sot] JSON parse failed on chunk {chunk_idx}: {err1}. Raw head: {snippet!r}", flush=True)
    return {"howtos": [], "insights": [], "concepts": [], "products": [], "_parse_error": str(err1)}


def _find_last_balanced_object_end(s):
    """Scan the string respecting JSON string literals; return the index of the outermost
    closing `}` that balances the first opening `{`. Returns None if no such balance found."""
    if not s:
        return None
    start = s.find("{")
    if start < 0:
        return None
    depth = 0
    in_str = False
    escape = False
    last_close = None
    for i in range(start, len(s)):
        ch = s[i]
        if in_str:
            if escape:
                escape = False
            elif ch == "\\":
                escape = True
            elif ch == '"':
                in_str = False
            continue
        if ch == '"':
            in_str = True
        elif ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                last_close = i
                # Don't break — keep scanning; outer object may close later.
                break
    return last_close


# ── Hallucination guards for how-tos ─────────────────────────────────────

_PROCEDURAL_PATTERNS = [
    re.compile(r"(?m)^\s*\d+[\.\)]\s+\S"),               # "1. Open Site Explorer" / "1) Open"
    re.compile(r"\bStep\s+\d+\b", re.IGNORECASE),         # "Step 1", "Step 2"
    re.compile(r"(?m)^\s*[-*]\s+\S+.{0,300}?\n\s*[-*]\s+\S"),  # consecutive bullets
    re.compile(r"\bfirst[,]?\s.*?\bthen\b.*?\b(next|finally|after that)\b", re.IGNORECASE | re.DOTALL),
    re.compile(r"\b(?:click|select|enter|navigate|paste|filter)\s+(?:the\s+)?[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*", ),  # imperative + named UI element
]


def _has_procedural_signals(text):
    """Does the chunk contain at least one procedural-language pattern? If not, skip how-to extraction."""
    if not text:
        return False
    hits = sum(1 for p in _PROCEDURAL_PATTERNS if p.search(text))
    return hits >= 1


def _normalize_for_quote_match(s):
    """Lowercase + collapse whitespace + strip surrounding punctuation. Used for substring match."""
    if not s:
        return ""
    s = s.lower()
    s = re.sub(r"[\s\u00a0\u200b]+", " ", s)        # collapse whitespace + nbsp + zero-width
    s = re.sub(r"[\u2018\u2019\u201c\u201d`]", "'", s)  # normalize quote chars
    s = re.sub(r"[\u2013\u2014]", "-", s)            # en/em dash to hyphen
    return s.strip()


def _validate_howto_quotes(howto, source_norm):
    """Mark each step with `verified` True/False based on whether `source_quote` appears in source.
    Mutates `howto` in place. Returns (verified_count, total_count).
    If a step has no source_quote field at all, treat as unverified."""
    steps = howto.get("steps") or []
    verified = 0
    norm_steps = []
    for s in steps:
        if isinstance(s, str):
            # Legacy string-only step — cannot verify.
            norm_steps.append({"text": s, "source_quote": "", "verified": False})
            continue
        text = s.get("text") or s.get("title") or ""
        quote = (s.get("source_quote") or "").strip()
        is_verified = False
        if quote:
            q_norm = _normalize_for_quote_match(quote)
            if q_norm and q_norm in source_norm:
                is_verified = True
        if is_verified:
            verified += 1
        norm_steps.append({"text": text, "source_quote": quote, "verified": is_verified})
    howto["steps"] = norm_steps
    return verified, len(steps)


def _extract_one_chunk(chunk_text, source_url, chunk_idx=1, total_chunks=1):
    """Run one LLM extraction pass on a single chunk. Returns (howtos, insights, concepts, products).
    Applies hallucination guards: skips how-to extraction entirely if no procedural signals;
    drops how-tos whose steps can't be quoted from source."""
    has_proc = _has_procedural_signals(chunk_text)
    hint_parts = []
    if source_url:
        hint_parts.append(f"Source URL (use for source_url unless the article cites a different canonical URL): {source_url}")
    if total_chunks > 1:
        hint_parts.append(f"This is chunk {chunk_idx} of {total_chunks} from a longer document. Extract only items fully visible in THIS chunk; don't speculate about what's in other chunks.")
    if not has_proc:
        hint_parts.append("NOTE: This chunk has no procedural language (no numbered lists, no 'Step N:', no sequential 'first/then'). Return [] for howtos no matter what product names are mentioned.")
    hint = ("\n" + "\n".join(hint_parts) + "\n") if hint_parts else ""
    resp = llm.chat.completions.create(
        model=LLM_MODEL,
        messages=[
            {"role": "system", "content": EXTRACT_SYSTEM},
            {"role": "user", "content": f"Extract Ahrefs how-tos, quotable insights, concept explainers, and product info from this content.{hint}\nCONTENT:\n{chunk_text}"}
        ],
        response_format={"type": "json_object"},
    )
    raw = resp.choices[0].message.content
    extracted = _parse_extraction_json(raw, chunk_idx=chunk_idx)
    howtos = extracted.get("howtos", []) or []
    insights = extracted.get("insights", []) or []
    concepts = extracted.get("concepts", []) or []
    products = extracted.get("products", []) or []

    # ── Validate how-to step quotes against the source text ──
    source_norm = _normalize_for_quote_match(chunk_text)
    validated = []
    for h in howtos:
        v, t = _validate_howto_quotes(h, source_norm)
        h["verified_steps"] = v
        h["total_steps"] = t
        # Drop the whole how-to if NONE of its steps could be quoted.
        if t > 0 and v == 0:
            print(f"[sot] dropped hallucinated how-to (0/{t} quotes match source): {h.get('title','')[:80]}", flush=True)
            continue
        validated.append(h)

    # Belt-and-suspenders: if the procedural gate said "no procedure here",
    # also drop any how-tos the model emitted (rare but the gate is the contract).
    if not has_proc and validated:
        print(f"[sot] gate dropped {len(validated)} how-to(s) from non-procedural chunk", flush=True)
        validated = []

    return validated, insights, concepts, products


def _dedupe_items(howtos, insights, concepts, products):
    """Merge near-duplicates that the chunker overlap might produce.
    Keeps the FIRST occurrence (chunks processed in order)."""
    def norm(s):
        return re.sub(r'[^a-z0-9]+', ' ', (s or '').lower()).strip()

    def _dedupe(items, key_field):
        seen = set()
        out = []
        for it in items:
            key = norm(it.get(key_field, ""))
            if not key or key in seen:
                continue
            seen.add(key)
            out.append(it)
        return out

    def _dedupe_products(items):
        # Composite key: product + claim. Same claim about different products = different fact.
        seen = set()
        out = []
        for it in items:
            key = norm((it.get("product") or "") + " :: " + (it.get("claim") or ""))
            if not key or key in seen:
                continue
            seen.add(key)
            out.append(it)
        return out

    return (_dedupe(howtos, "title"),
            _dedupe(insights, "claim"),
            _dedupe(concepts, "concept"),
            _dedupe_products(products))


def _extract_one(actual_content, source_url, progress_cb=None):
    """Run extraction over content, chunking if needed.
    Returns (howtos, insights, concepts, products, chunk_count)."""
    chunks = _chunk_text(actual_content)
    total_chunks = len(chunks)
    all_h, all_i, all_c, all_p = [], [], [], []
    for idx, total, chunk in chunks:
        if progress_cb:
            try: progress_cb(idx, total)
            except Exception: pass
        h, i, c, p = _extract_one_chunk(chunk, source_url, idx, total)
        all_h.extend(h)
        all_i.extend(i)
        all_c.extend(c)
        all_p.extend(p)
    if total_chunks > 1:
        all_h, all_i, all_c, all_p = _dedupe_items(all_h, all_i, all_c, all_p)
    return all_h, all_i, all_c, all_p, total_chunks


def _is_url_safe(url):
    """SSRF guard: only allow public http(s) URLs.

    Rejects:
      - non-http(s) schemes (file://, gopher://, ftp://, data:, etc.)
      - hosts that resolve to loopback / private / link-local / reserved IPs
      - bare IP literals in private ranges
    """
    from urllib.parse import urlparse
    import socket
    import ipaddress

    try:
        parsed = urlparse(url)
    except Exception:
        return False, "unparseable URL"
    if parsed.scheme not in ("http", "https"):
        return False, f"scheme {parsed.scheme!r} not allowed"
    host = parsed.hostname
    if not host:
        return False, "missing host"
    try:
        infos = socket.getaddrinfo(host, None)
    except socket.gaierror:
        return False, "DNS lookup failed"
    for info in infos:
        addr = info[4][0]
        try:
            ip = ipaddress.ip_address(addr)
        except ValueError:
            continue
        if (
            ip.is_private
            or ip.is_loopback
            or ip.is_link_local
            or ip.is_reserved
            or ip.is_multicast
            or ip.is_unspecified
        ):
            return False, f"host resolves to non-public address {addr}"
    return True, None


def _fetch_url(url):
    """Fetch a URL and return reasonably clean text.

    Uses `trafilatura` if available (best extraction), otherwise falls back
    to `requests` + a minimal HTML strip. Returns (text, error).
    """
    ok, why = _is_url_safe(url)
    if not ok:
        return None, f"refused: {why}"
    try:
        import requests
    except ImportError:
        return None, "requests not installed"
    try:
        resp = requests.get(
            url,
            timeout=30,
            allow_redirects=False,  # don't let a 302 sneak past _is_url_safe
            headers={"User-Agent": "Mozilla/5.0 (SourceOfTruth/1.0)"},
        )
        resp.raise_for_status()
    except Exception as e:
        return None, f"fetch failed: {e}"[:200]

    html = resp.text
    # Prefer trafilatura for clean text extraction
    try:
        import trafilatura  # type: ignore
        text = trafilatura.extract(html, include_comments=False, include_tables=True)
        if text:
            return text[:FETCH_MAX_LENGTH], None
    except ImportError:
        pass

    # Fallback: minimal strip
    import re as _re
    text = _re.sub(r"<script[\s\S]*?</script>", " ", html, flags=_re.I)
    text = _re.sub(r"<style[\s\S]*?</style>", " ", text, flags=_re.I)
    text = _re.sub(r"<[^>]+>", " ", text)
    text = _re.sub(r"\s+", " ", text).strip()
    return text[:FETCH_MAX_LENGTH], None


def _extract_text_from_upload(filename, raw_bytes):
    """Extract plain text from a .md / .docx / .pdf upload. Returns (text, error)."""
    name = (filename or "").lower()
    try:
        if name.endswith(".md") or name.endswith(".markdown") or name.endswith(".txt"):
            text = raw_bytes.decode("utf-8", errors="replace")
            return text, None

        if name.endswith(".docx"):
            from io import BytesIO
            from docx import Document
            doc = Document(BytesIO(raw_bytes))
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            # Tables
            for tbl in doc.tables:
                for row in tbl.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts), None

        if name.endswith(".pdf"):
            import pdfplumber
            from io import BytesIO
            parts = []
            with pdfplumber.open(BytesIO(raw_bytes)) as pdf:
                for page in pdf.pages:
                    t = page.extract_text() or ""
                    if t.strip():
                        parts.append(t)
            return "\n\n".join(parts), None

        return None, f"Unsupported file type: {name}. Use .md, .docx, or .pdf."
    except Exception as e:
        return None, f"Parse error: {type(e).__name__}: {str(e)[:200]}"



def _tag_items(howtos, insights, concepts, products, source_url, today):
    """Stamp source_url + date defaults on all four buckets so the review UI can group by source."""
    for h in howtos:
        if source_url:
            h.setdefault("source_url", source_url)
    for ins in insights:
        if source_url and not ins.get("source_url"):
            ins["source_url"] = source_url
        if not ins.get("date"):
            ins["date"] = today
    for c in concepts:
        if source_url and not c.get("source_url"):
            c["source_url"] = source_url
        if not c.get("date"):
            c["date"] = today
    for p in products:
        if source_url and not p.get("source_url"):
            p["source_url"] = source_url
        if not p.get("date"):
            p["date"] = today


def _attach_lookalikes(proposed):
    """Run duplicate-detection against existing GitHub content for each bucket.
    Attaches three aligned arrays: lookalikes_howtos / _insights / _concepts.
    Each entry is None (new) or {candidates, default_action, best_match}."""
    import traceback
    for bucket, key in [("insights", "lookalikes_insights"), ("concepts", "lookalikes_concepts"), ("howtos", "lookalikes_howtos"), ("products", "lookalikes_products")]:
        items = proposed.get(bucket, [])
        try:
            res = find_lookalikes(bucket, items)
            matched = sum(1 for r in res if r and r.get("best_match"))
            print(f"[sot] {bucket} look-alikes: {matched}/{len(items)} matched", flush=True)
            proposed[key] = res
        except Exception as e:
            print(f"[sot] {bucket} lookalike scan failed: {e}\n{traceback.format_exc()}", flush=True)
            proposed[key] = [None] * len(items)


@blueprint.route("/api/ingest", methods=["POST"])
def api_ingest():
    from datetime import date
    data = request.get_json() or {}
    source_type = data.get("type", "text")
    content = (data.get("content") or "").strip()
    urls = data.get("urls") or []

    if source_type == "bulk_urls":
        # Normalize: strip + drop empties + dedupe (preserve order)
        seen = set()
        clean_urls = []
        for u in urls:
            u = (u or "").strip()
            if not u or u in seen:
                continue
            seen.add(u)
            clean_urls.append(u)
        if not clean_urls:
            return jsonify({"error": "No URLs provided"}), 400
        if len(clean_urls) > 10:
            return jsonify({"error": "Maximum 10 URLs per batch"}), 400
        stored_content = json.dumps(clean_urls)
        source_ref = f"{len(clean_urls)} URLs"
    else:
        if not content:
            return jsonify({"error": "No content provided"}), 400
        stored_content = content
        source_ref = content[:200] if source_type == "url" else None

    conn = get_db()
    cur = conn.cursor()
    cur.execute("""INSERT INTO sot_pending_updates (source_type, source_ref, source_content, status)
                   VALUES (%s, %s, %s, 'extracting') RETURNING id""",
                 (source_type, source_ref, stored_content))
    pu_id = cur.fetchone()[0]
    conn.commit()
    conn.close()

    job_id = str(uuid.uuid4())
    with _jobs_lock:
        _jobs[job_id] = {"status": "running", "pending_id": pu_id, "phase": "starting"}

    def run_extraction():
        try:
            today = date.today().isoformat()
            all_howtos = []
            all_insights = []
            all_concepts = []
            all_products = []
            per_url_status = []

            def _make_progress_cb(prefix):
                def cb(chunk_idx, total):
                    if total > 1:
                        with _jobs_lock:
                            _jobs[job_id]["phase"] = f"{prefix} — chunk {chunk_idx}/{total}"
                return cb

            if source_type == "bulk_urls":
                url_list = clean_urls
                for idx, url in enumerate(url_list):
                    prefix = f"{idx+1}/{len(url_list)}: {url[:60]}"
                    with _jobs_lock:
                        _jobs[job_id]["phase"] = prefix
                    fetched, ferr = _fetch_url(url)
                    if ferr or not fetched:
                        per_url_status.append({"url": url, "status": "fetch_error", "error": ferr or "empty content"})
                        continue
                    try:
                        howtos, insights, concepts, products, n_chunks = _extract_one(fetched, url, progress_cb=_make_progress_cb(prefix))
                    except Exception as ex:
                        per_url_status.append({"url": url, "status": "extract_error", "error": str(ex)[:200]})
                        continue
                    _tag_items(howtos, insights, concepts, products, url, today)
                    all_howtos.extend(howtos)
                    all_insights.extend(insights)
                    all_concepts.extend(concepts)
                    all_products.extend(products)
                    per_url_status.append({"url": url, "status": "ok", "howto_count": len(howtos), "insight_count": len(insights), "concept_count": len(concepts), "product_count": len(products), "chunks": n_chunks})
            else:
                actual_content = stored_content
                source_url = stored_content if source_type == "url" else ""
                if source_type == "url":
                    fetched, ferr = _fetch_url(stored_content)
                    actual_content = fetched if fetched else f"[Failed to fetch URL: {stored_content}]\n{ferr or ''}"
                with _jobs_lock:
                    _jobs[job_id]["phase"] = "extracting"
                howtos, insights, concepts, products, _ = _extract_one(actual_content, source_url, progress_cb=_make_progress_cb("extracting"))
                _tag_items(howtos, insights, concepts, products, source_url, today)
                all_howtos = howtos
                all_insights = insights
                all_concepts = concepts
                all_products = products

            proposed = {"howtos": all_howtos, "insights": all_insights, "concepts": all_concepts, "products": all_products, "per_url": per_url_status}

            with _jobs_lock:
                _jobs[job_id]["phase"] = "checking for look-alikes…"
            _attach_lookalikes(proposed)

            conn2 = get_db()
            conn2.cursor().execute("""UPDATE sot_pending_updates
                SET extracted_facts = %s, proposed_changes = %s, status = 'pending' WHERE id = %s""",
                (json.dumps({"howtos_count": len(all_howtos), "insights_count": len(all_insights), "concepts_count": len(all_concepts), "products_count": len(all_products)}), json.dumps(proposed), pu_id))
            conn2.commit()
            conn2.close()

            with _jobs_lock:
                _jobs[job_id] = {"status": "done", "pending_id": pu_id}

        except Exception as e:
            conn3 = get_db()
            conn3.cursor().execute("UPDATE sot_pending_updates SET status = 'error' WHERE id = %s", (pu_id,))
            conn3.commit()
            conn3.close()
            with _jobs_lock:
                _jobs[job_id] = {"status": "error", "error": str(e), "pending_id": pu_id}

    threading.Thread(target=run_extraction, daemon=True).start()
    return jsonify({"job_id": job_id, "pending_id": pu_id})


@blueprint.route("/api/ingest/files", methods=["POST"])
def api_ingest_files():
    from datetime import date
    files = request.files.getlist("files")
    note = (request.form.get("note") or "").strip()

    if not files:
        return jsonify({"error": "No files uploaded"}), 400
    if len(files) > 10:
        return jsonify({"error": "Maximum 10 files per batch"}), 400

    # Read all file bytes up front (multipart parts don't survive into the background thread cleanly)
    file_inputs = []
    for f in files:
        if not f or not f.filename:
            continue
        name = f.filename
        raw = f.read()
        if not raw:
            continue
        file_inputs.append({"name": name, "bytes": raw, "size": len(raw)})

    if not file_inputs:
        return jsonify({"error": "All uploads were empty"}), 400

    stored_meta = {
        "files": [{"name": fi["name"], "size": fi["size"]} for fi in file_inputs],
        "note": note,
    }
    source_ref = f"{len(file_inputs)} file{'s' if len(file_inputs) != 1 else ''}" + (f" · {note[:80]}" if note else "")

    conn = get_db()
    cur = conn.cursor()
    cur.execute("""INSERT INTO sot_pending_updates (source_type, source_ref, source_content, status)
                   VALUES (%s, %s, %s, 'extracting') RETURNING id""",
                 ("files", source_ref, json.dumps(stored_meta)))
    pu_id = cur.fetchone()[0]
    conn.commit()
    conn.close()

    job_id = str(uuid.uuid4())
    with _jobs_lock:
        _jobs[job_id] = {"status": "running", "pending_id": pu_id, "phase": "starting"}

    def run_extraction():
        try:
            today = date.today().isoformat()
            all_howtos = []
            all_insights = []
            all_concepts = []
            all_products = []
            per_file_status = []

            def _make_file_progress_cb(prefix):
                def cb(chunk_idx, total):
                    if total > 1:
                        with _jobs_lock:
                            _jobs[job_id]["phase"] = f"{prefix} — chunk {chunk_idx}/{total}"
                return cb

            for idx, fi in enumerate(file_inputs):
                prefix = f"{idx+1}/{len(file_inputs)}: {fi['name'][:60]}"
                with _jobs_lock:
                    _jobs[job_id]["phase"] = prefix
                text, err = _extract_text_from_upload(fi["name"], fi["bytes"])
                if err or not text or not text.strip():
                    per_file_status.append({"name": fi["name"], "status": "parse_error", "error": err or "empty content"})
                    continue
                tag = fi["name"]
                try:
                    howtos, insights, concepts, products, n_chunks = _extract_one(text, tag, progress_cb=_make_file_progress_cb(prefix))
                except Exception as ex:
                    per_file_status.append({"name": fi["name"], "status": "extract_error", "error": str(ex)[:200]})
                    continue
                # Override source_url with filename + attach note as source_title
                for h in howtos:
                    h["source_url"] = tag
                    if note:
                        h["source_title"] = note
                    h["is_file"] = True
                for ins in insights:
                    ins["source_url"] = tag
                    if note and not ins.get("source_title"):
                        ins["source_title"] = note
                    ins["is_file"] = True
                    if not ins.get("date"):
                        ins["date"] = today
                for c in concepts:
                    c["source_url"] = tag
                    if note and not c.get("source_title"):
                        c["source_title"] = note
                    c["is_file"] = True
                    if not c.get("date"):
                        c["date"] = today
                for p in products:
                    p["source_url"] = tag
                    if note and not p.get("source_title"):
                        p["source_title"] = note
                    p["is_file"] = True
                    if not p.get("date"):
                        p["date"] = today
                all_howtos.extend(howtos)
                all_insights.extend(insights)
                all_concepts.extend(concepts)
                all_products.extend(products)
                per_file_status.append({"name": fi["name"], "status": "ok",
                                        "howto_count": len(howtos), "insight_count": len(insights),
                                        "concept_count": len(concepts), "product_count": len(products),
                                        "text_length": len(text), "chunks": n_chunks})

            proposed = {"howtos": all_howtos, "insights": all_insights, "concepts": all_concepts, "products": all_products, "per_file": per_file_status, "note": note}

            with _jobs_lock:
                _jobs[job_id]["phase"] = "checking for look-alikes…"
            _attach_lookalikes(proposed)

            conn2 = get_db()
            conn2.cursor().execute("""UPDATE sot_pending_updates
                SET extracted_facts = %s, proposed_changes = %s, status = 'pending' WHERE id = %s""",
                (json.dumps({"howtos_count": len(all_howtos), "insights_count": len(all_insights), "concepts_count": len(all_concepts), "products_count": len(all_products)}),
                 json.dumps(proposed), pu_id))
            conn2.commit()
            conn2.close()

            with _jobs_lock:
                _jobs[job_id] = {"status": "done", "pending_id": pu_id}

        except Exception as e:
            conn3 = get_db()
            conn3.cursor().execute("UPDATE sot_pending_updates SET status = 'error' WHERE id = %s", (pu_id,))
            conn3.commit()
            conn3.close()
            with _jobs_lock:
                _jobs[job_id] = {"status": "error", "error": str(e), "pending_id": pu_id}

    threading.Thread(target=run_extraction, daemon=True).start()
    return jsonify({"job_id": job_id, "pending_id": pu_id})


@blueprint.route("/api/ingest/status/<job_id>")
def ingest_status(job_id):
    with _jobs_lock:
        job = _jobs.get(job_id)
    if not job:
        return jsonify({"error": "Not found"}), 404
    return jsonify(job)


@blueprint.route("/review/<int:pu_id>")
def review(pu_id):
    conn = get_db()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    cur.execute("SELECT * FROM sot_pending_updates WHERE id = %s", (pu_id,))
    pu = cur.fetchone()
    conn.close()
    if not pu:
        return redirect(url_for("source_of_truth.dashboard"))
    return render_template("sot_review.html", pu=pu)


@blueprint.route("/api/review/<int:pu_id>/apply", methods=["POST"])
def apply_changes(pu_id):
    """Apply per-item actions for each bucket.
    Body: {
      actions: {
        insights: {"<idx>": "accept_new"|"skip"|"merge:<card_id>"},
        concepts: {...},
        howtos:   {...},
      }
    }
    Falls back to legacy accepted_* arrays for backward compatibility.
    """
    data = request.get_json() or {}
    actions = data.get("actions") or {}

    conn = get_db()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    cur.execute("SELECT * FROM sot_pending_updates WHERE id = %s", (pu_id,))
    pu = cur.fetchone()
    if not pu or pu['status'] != 'pending':
        conn.close()
        return jsonify({"error": "Invalid or already processed"}), 400

    proposed = pu['proposed_changes'] or {}
    howtos = proposed.get("howtos", [])
    insights = proposed.get("insights", [])
    concepts = proposed.get("concepts", [])
    products = proposed.get("products", [])

    # Backward compat: convert old accepted_* arrays to actions
    if not actions:
        actions = {"howtos": {}, "insights": {}, "concepts": {}, "products": {}}
        for i in data.get("accepted_howtos", []):
            actions["howtos"][str(i)] = "accept_new"
        for i in data.get("accepted_insights", []):
            actions["insights"][str(i)] = "accept_new"
        for i in data.get("accepted_concepts", []):
            actions["concepts"][str(i)] = "accept_new"
        for i in data.get("accepted_products", []):
            actions["products"][str(i)] = "accept_new"

    summary = {"howtos_new": 0, "howtos_skipped": 0,
               "insights_new": 0, "insights_merged": 0, "insights_skipped": 0,
               "concepts_new": 0, "concepts_skipped": 0,
               "products_new": 0, "products_merged": 0, "products_skipped": 0}

    # ── How-tos ──
    new_howtos = []
    for idx_str, act in (actions.get("howtos") or {}).items():
        idx = int(idx_str)
        if idx >= len(howtos):
            continue
        if act == "accept_new":
            new_howtos.append(howtos[idx])
        else:
            summary["howtos_skipped"] += 1
    for h in new_howtos:
        normalized_steps = []
        for s in h.get("steps", []):
            if isinstance(s, str):
                normalized_steps.append({"title": s, "description": ""})
            else:
                title = s.get("text") or s.get("title") or ""
                # Tag unverified steps in the saved how-to so the user can spot them later.
                if s.get("verified") is False:
                    title = f"{title} [unverified]"
                normalized_steps.append({"title": title, "description": ""})
        save_howto(h["title"], "", "", normalized_steps, [])
        summary["howtos_new"] += 1
        time.sleep(1.2)

    # ── Concepts ──
    new_concepts = []
    for idx_str, act in (actions.get("concepts") or {}).items():
        idx = int(idx_str)
        if idx >= len(concepts):
            continue
        if act == "accept_new":
            new_concepts.append(concepts[idx])
        else:
            summary["concepts_skipped"] += 1
    if new_concepts:
        append_concepts(new_concepts)
        summary["concepts_new"] = len(new_concepts)
        time.sleep(1.5)

    # ── Insights (merge supported) ──
    new_insights = []
    merges = []  # list of (target_card_id, item)
    for idx_str, act in (actions.get("insights") or {}).items():
        idx = int(idx_str)
        if idx >= len(insights):
            continue
        if act == "accept_new":
            new_insights.append(insights[idx])
        elif act.startswith("merge:"):
            merges.append((act.split(":", 1)[1], insights[idx]))
        else:
            summary["insights_skipped"] += 1

    # Do merges sequentially BEFORE appending new ones (cleaner audit trail)
    for target_id, item in merges:
        ok, err = merge_source_into_insight(
            target_id,
            item.get("source_url", ""),
            item.get("source_title"),
        )
        if ok:
            summary["insights_merged"] += 1
        else:
            print(f"[sot] merge failed for {target_id}: {err}")
            # Fall back to accepting as new
            new_insights.append(item)
        time.sleep(1.2)

    if new_insights:
        append_insights(new_insights)
        summary["insights_new"] = len(new_insights)

    # ── Products (merge supported, same shape as insights) ──
    new_products = []
    product_merges = []
    for idx_str, act in (actions.get("products") or {}).items():
        idx = int(idx_str)
        if idx >= len(products):
            continue
        if act == "accept_new":
            new_products.append(products[idx])
        elif act.startswith("merge:"):
            product_merges.append((act.split(":", 1)[1], products[idx]))
        else:
            summary["products_skipped"] += 1

    for target_id, item in product_merges:
        ok, err = merge_source_into_product(
            target_id,
            item.get("source_url", ""),
            item.get("source_title"),
        )
        if ok:
            summary["products_merged"] += 1
        else:
            print(f"[sot] product merge failed for {target_id}: {err}")
            new_products.append(item)
        time.sleep(1.2)

    if new_products:
        append_products(new_products)
        summary["products_new"] = len(new_products)

    cur2 = conn.cursor()
    cur2.execute("UPDATE sot_pending_updates SET status = 'applied' WHERE id = %s", (pu_id,))
    conn.commit()
    conn.close()
    return jsonify({"ok": True, "summary": summary})


@blueprint.route("/api/review/<int:pu_id>/dismiss", methods=["POST"])
def dismiss_changes(pu_id):
    conn = get_db()
    cur = conn.cursor()
    cur.execute("UPDATE sot_pending_updates SET status = 'dismissed' WHERE id = %s", (pu_id,))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})
